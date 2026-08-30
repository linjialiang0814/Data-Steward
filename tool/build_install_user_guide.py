"""Build the standalone Data Steward installation and user guide DOCX.

The source of truth is docs/INSTALL_AND_USER_GUIDE.md. The generated DOCX
embeds all screenshots so it can be distributed without the repository.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x20, 0x37, 0x48)
MUTED = RGBColor(0x5F, 0x6B, 0x78)
GOLD = RGBColor(0xA6, 0x73, 0x19)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    east_asia: str = "Microsoft YaHei",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    props = run._element.get_or_add_rPr()
    fonts = props.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        props.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_paragraph_bottom_border(paragraph, color="D7DBE2", size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(3)
    left = hp.add_run("DATA STEWARD")
    set_run_font(left, size=9, color=DARK_BLUE, bold=True)
    right = hp.add_run("   |   安装与使用指南")
    set_run_font(right, size=9, color=MUTED)
    set_paragraph_bottom_border(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(fp)

    doc.core_properties.title = "Data Steward 安装与使用指南"
    doc.core_properties.subject = "Windows 与 Android APP Demo 安装、配置和使用"
    doc.core_properties.creator = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.keywords = "Data Steward, Windows, Android, Hermes, 安装指南"


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(116)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("APP DEMO · OPERATOR GUIDE")
    set_run_font(run, size=10.5, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Data Steward")
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run("多设备数据管理智能管家")
    set_run_font(run, size=16, color=DARK_BLUE, bold=True)

    guide = doc.add_paragraph()
    guide.alignment = WD_ALIGN_PARAGRAPH.CENTER
    guide.paragraph_format.space_after = Pt(30)
    run = guide.add_run("安装与使用指南")
    set_run_font(run, size=14, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(4)
    run = meta.add_run("Windows 11 + Android · APP Demo 1.0.0+1 · S6-G")
    set_run_font(run, size=10.5, color=NAVY, bold=True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("安全配对 · 跨设备会话 · Hermes 智能工作流 · 确认式文件操作")
    set_run_font(run, size=9.5, color=MUTED, italic=True)
    doc.add_page_break()


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text: str, *, default_size: float = 11) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=default_size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(
                run,
                name="Consolas",
                east_asia="Microsoft YaHei",
                size=default_size - 0.5,
                color=DARK_BLUE,
            )
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=default_size)


def add_callout(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    add_inline(p, text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    for index, line in enumerate(lines):
        if index:
            run = p.add_run()
            run.add_break()
        run = p.add_run(line or " ")
        set_run_font(
            run,
            name="Consolas",
            east_asia="Microsoft YaHei",
            size=8.5,
            color=NAVY,
        )
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_screenshot(doc: Document, source: Path, alt: str) -> None:
    if not source.is_file():
        raise FileNotFoundError(source)
    with Image.open(source) as image:
        width, height = image.size
    portrait = height / width > 1.25
    display_width = 2.65 if portrait else 5.95
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(source), width=Inches(display_width))
    picture = p._p.xpath(".//wp:docPr")
    if picture:
        picture[0].set("descr", alt)


def build(source: Path, output: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    lines = source.read_text(encoding="utf-8").splitlines()
    repo_root = source.parent.parent
    in_code = False
    code_lines: list[str] = []
    skip_intro_metadata = True
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        p = doc.add_paragraph()
        add_inline(p, " ".join(part.strip() for part in paragraph_buffer))
        paragraph_buffer = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            flush_paragraph()
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line == "<!-- pagebreak -->":
            flush_paragraph()
            doc.add_page_break()
            continue
        if line.startswith("# "):
            continue
        if skip_intro_metadata and line.startswith("> "):
            continue
        if line and not line.startswith("> "):
            skip_intro_metadata = False
        image_match = re.fullmatch(r"!\[(.+?)\]\((.+?)\)", line)
        if image_match:
            flush_paragraph()
            add_screenshot(
                doc,
                (source.parent / image_match.group(2)).resolve(),
                image_match.group(1),
            )
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            level = min(len(heading_match.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading_match.group(2), default_size={1: 16, 2: 13, 3: 12}[level])
            continue
        if line.startswith("> "):
            flush_paragraph()
            add_callout(doc, line[2:])
            continue
        bullet_match = re.match(r"^-\s+(.+)$", line)
        if bullet_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, bullet_match.group(1))
            continue
        number_match = re.match(r"^\d+\.\s+(.+)$", line)
        if number_match:
            flush_paragraph()
            p = doc.add_paragraph(style="List Number")
            add_inline(p, number_match.group(1))
            continue
        if line.startswith("图 ") or re.match(r"^图\s*\d+", line):
            flush_paragraph()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = False
            run = p.add_run(line)
            set_run_font(run, size=9, color=MUTED, italic=True)
            continue
        if not line:
            flush_paragraph()
            continue
        paragraph_buffer.append(line)

    flush_paragraph()
    if in_code:
        add_code_block(doc, code_lines)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--source",
        type=Path,
        default=Path("docs/INSTALL_AND_USER_GUIDE.md"),
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("dist/DataSteward-安装与使用指南.docx"),
    )
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve())
    print(args.output.resolve())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
